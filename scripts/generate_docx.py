#!/usr/bin/env python3
"""Generate 架构设计文档.docx from the architecture plan."""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page Setup ──
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── Style Setup ──
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.5

# Fix CJK font
rPr = style.element.get_or_add_rPr()
rFonts = OxmlElement('w:rFonts')
rFonts.set(qn('w:eastAsia'), '微软雅黑')
rPr.insert(0, rFonts)

# Heading styles
for level, (size, color_hex) in enumerate([
    (22, '1a1a2e'),  # Heading 1
    (16, '2d3436'),  # Heading 2
    (13, '636e72'),  # Heading 3
], start=1):
    h_style = doc.styles[f'Heading {level}']
    h_font = h_style.font
    h_font.name = '微软雅黑'
    h_font.size = Pt(size)
    h_font.bold = True
    h_font.color.rgb = RGBColor.from_string(color_hex)
    h_style.paragraph_format.space_before = Pt(18 if level == 1 else 14)
    h_style.paragraph_format.space_after = Pt(10)
    h_style.paragraph_format.line_spacing = 1.2
    hRpr = h_style.element.get_or_add_rPr()
    hFonts = OxmlElement('w:rFonts')
    hFonts.set(qn('w:eastAsia'), '微软雅黑')
    hRpr.insert(0, hFonts)

# Code style
code_style = doc.styles.add_style('CodeBlock', WD_STYLE_TYPE.PARAGRAPH)
code_font = code_style.font
code_font.name = 'Consolas'
code_font.size = Pt(9)
code_style.paragraph_format.space_before = Pt(4)
code_style.paragraph_format.space_after = Pt(4)
code_style.paragraph_format.line_spacing = 1.15

# ── Helper Functions ──

def add_title_page():
    """Add title page."""
    # Empty lines for spacing
    for _ in range(4):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('糖糖桌宠 架构设计文档')
    run.font.size = Pt(32)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string('1a1a2e')
    run.font.name = '微软雅黑'
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), '微软雅黑')
    rPr.insert(0, rFonts)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Desk Pet — Architecture Design Document')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor.from_string('636e72')
    run.font.italic = True

    doc.add_paragraph()
    doc.add_paragraph()

    ver = doc.add_paragraph()
    ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = ver.add_run('版本: v2.0    日期: ' + datetime.date.today().strftime('%Y-%m-%d'))
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string('636e72')

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run('双模式桌面宠物助手架构')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string('636e72')

    doc.add_page_break()


def add_code(text, indent=0):
    """Add a code block with monospace font."""
    for line in text.strip().split('\n'):
        p = doc.add_paragraph(style='CodeBlock')
        if indent:
            p.paragraph_format.left_indent = Cm(indent)
        p.add_run(line)


def add_diagram_box(title, text):
    """Add a bordered box containing a diagram."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(title)
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string('546e7a')

    for line in text.strip().split('\n'):
        p = doc.add_paragraph(style='CodeBlock')
        p.add_run(line)


def add_table(headers, rows):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.name = '微软雅黑'

    # Data rows
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.name = '微软雅黑'

    doc.add_paragraph()  # spacing after table


def add_body(text):
    """Add body text paragraph."""
    p = doc.add_paragraph(text)
    return p

# ── Document Content ──

add_title_page()

# ===================
# 第一章
# ===================
doc.add_heading('一、设计理念与定位', level=1)

add_body('糖糖桌宠是一个桌面宠物助手应用。像素角色常驻桌面，以无边框透明窗口置顶显示。用户感知为"和桌宠对话"，底层运行一个完整的 AI Agent。')

add_body('核心原则：')
add_body('1. 宠物是界面，Agent 是引擎 — 用户永远感知为宠物对话，底层 Agent 透明工作')
add_body('2. 双模式，同一套 Loop — 轻量/助手模式共用核心引擎，仅工具可见范围不同')
add_body('3. 轻量不残废 — 默认模式具备文件读取、目录浏览、系统信息、安全命令等基本能力')
add_body('4. 人格中间件 — 工具调用、结果、错误全部经人格层转换为角色化表达')
add_body('5. 安全内建 — 不在各处分散判断，统一入口强制校验')
add_body('6. 思考强度可配 — auto/low/medium/high，auto 根据任务复杂度自动调节')

add_diagram_box('交互示意', '''
┌──────────────────────────────────────────────┐
│              用户看到的                        │
│   "糖糖，帮我看看桌面有什么文件"                │
│   "好的主人～让我看看...有3个文件夹和2个文档～" │
│                                              │
│              实际发生的                        │
│   PreProcessor → SystemPrompt(人格+记忆)       │
│   → AI(工具调用: ls ~/Desktop)                 │
│   → SafetyControl(通过) → bash.exec           │
│   → 结果回注 → AI(生成自然语言)                 │
│   → 人格过滤 → UI展示                          │
└──────────────────────────────────────────────┘
''')

# ===================
# 第二章
# ===================
doc.add_heading('二、双模式架构', level=1)

doc.add_heading('2.1 设计思路', level=2)
add_body('不是两套系统，而是同一套 Agent Loop + 不同的工具可见范围。轻量模式和助手模式共用 PreProcessor、AgentLoop、ReplyGenerator 等全部核心引擎组件，唯一区别在于 ToolRegistry 中注册的工具数量和 Safety 的严格程度。')

add_diagram_box('双模式架构概览', '''
                    同一套 Agent Loop（始终运行）
                    ┌─────────────────────────────┐
                    │  PreProcessor                │
                    │  → ContextEngine.build()     │
                    │  → AI request                │
                    │  → Parse output              │
                    │  → Tool calls? → Execute     │
                    │  → ReplyGenerator            │
                    └─────────────┬───────────────┘
                                  │
                    ┌─────────────┴───────────────┐
                    │                             │
              轻量模式(默认)                  助手模式(设置中开启)
              │                             │
              ToolRegistry 注册:              ToolRegistry 额外注册:
              ├── file.read (SAFE)           ├── file.write (DANGER)
              ├── file.list (SAFE)           ├── bash 全量 (DANGER)
              ├── file.search (SAFE)         ├── app.open (NORMAL)
              ├── system.info (SAFE)         ├── clipboard (NORMAL/DANGER)
              ├── bash 白名单 (NORMAL)        ├── MCP servers (动态)
              ├── http.get (NORMAL)          ├── Skills (全部)
              └── memory.* (内部API)          └── agent.spawn (NORMAL)
              │                             │
              Safety: 简洁拒绝               Safety: 完整四级+确认UI
              Plan: 不启用                   Plan: 复杂任务时启用
              MCP/Skill: 不加载               MCP/Skill: 完整加载
              内存 +0 MB                     内存 +3-7 MB (+MCP子进程)
''')

doc.add_heading('2.2 能力对比', level=2)

add_table(
    ['能力', '轻量模式', '助手模式'],
    [
        ['AI 聊天 + 人格系统', '✅', '✅'],
        ['窗口感知主动搭话', '✅', '✅'],
        ['表情 / 音效 / 人格中间件', '✅', '✅'],
        ['记忆系统 (短期 + 长期)', '✅', '✅'],
        ['记忆整理', '✅ (定时器)', '✅ (fork agent)'],
        ['读文件 file.read', '✅', '✅'],
        ['列目录 file.list', '✅', '✅'],
        ['搜索文件 file.search', '✅', '✅'],
        ['系统信息 system.info', '✅', '✅'],
        ['安全命令 (ls/cat/find...)', '✅', '✅'],
        ['网页获取 http.get', '✅', '✅'],
        ['写文件 file.write', '❌', '✅'],
        ['全量 Bash (含危险命令)', '❌', '✅'],
        ['打开应用 app.open', '❌', '✅'],
        ['剪贴板操作', '❌', '✅'],
        ['MCP Server (外部工具)', '❌', '✅'],
        ['Skill (工具编排)', '❌', '✅'],
        ['SubAgent (fork/team)', '❌', '✅'],
        ['Plan 步骤 (复杂规划)', '❌', '✅'],
        ['完整安全确认 UI', '❌ (简洁拒绝)', '✅'],
    ]
)

doc.add_heading('2.3 模式切换', level=2)
add_body('配置项 assistantMode 控制模式（默认 false = 轻量模式）。切换时：')

add_body('轻量 → 助手：')
add_body('  1. dynamic import 加载 local-extra/ mcp/ skill/ safety/full-checker 模块')
add_body('  2. 追加注册助手专属工具到 ToolRegistry')
add_body('  3. 初始化 MCP Manager（异步连接，不阻塞 UI）')
add_body('  4. 加载 Skill 目录，Safety 升级为完整模式')

add_body('助手 → 轻量：')
add_body('  1. 断开所有 MCP 连接，kill stdio 子进程')
add_body('  2. 从 ToolRegistry 移除助手专属工具')
add_body('  3. 释放模块引用，让 GC 回收内存')
add_body('  4. Safety 降级为简洁模式')

# ===================
# 第三章
# ===================
doc.add_heading('三、模型思考强度', level=1)

doc.add_heading('3.1 概念', level=2)
add_body('思考强度是与双模式正交的独立维度，控制 AI 推理深度。配置项 thinkingEffort 支持四个级别：')

add_table(
    ['强度', '说明', '适用场景', '感知延迟', 'Token 消耗'],
    [
        ['low', '快速响应，轻推理', '闲聊 / 打招呼 / 简单表情', '< 1s', '低'],
        ['medium', '均衡', '日常对话 / 简单工具调用', '1-3s', '中'],
        ['high', '深度推理', '复杂分析 / 多步工具链', '3-10s', '高'],
        ['auto', '自动调节 (默认)', '根据任务复杂度自动切换', '动态', '动态'],
    ]
)

doc.add_heading('3.2 auto 模式自动切换逻辑', level=2)
add_body('在 auto 模式下，系统根据任务特征自动选择思考强度：')
add_body('• 闲聊 / 表情 / 打招呼 → low (快速响应)')
add_body('• "帮我看看" / "查一下" / "找一下" → medium (单步工具)')
add_body('• 工具调用 ≥ 2 轮 / Plan 触发 / "分析" "整理" "总结" → high (多步推理)')
add_body('• 窗口主动搭话 → low (简单评论即可)')
add_body('• 错误恢复 / 重试 → high (需要更仔细推理)')

doc.add_heading('3.3 配置与实现', level=2)
add_table(
    ['Provider', '实现方式'],
    [
        ['OpenAI (o-series)', 'reasoning_effort 参数'],
        ['Anthropic', 'thinking.budget_tokens 参数'],
        ['DeepSeek', 'reasoning_effort 或切换 deepseek-reasoner 模型'],
        ['通用 OpenAI 兼容', 'SystemPrompt 追加 "请简要回答" / "请仔细思考后回答"'],
        ['Ollama', 'num_predict + SystemPrompt 引导'],
    ]
)

# ===================
# 第四章
# ===================
doc.add_heading('四、系统分层架构', level=1)

add_diagram_box('分层架构图', '''
┌──────────────────────────────────────────────────────────┐
│  UI 层 (Vue 3)                                            │
│  StreamView(角色动画) / ChatPanel(对话) / TitleBar / Settings │
├──────────────────────────────────────────────────────────┤
│  宠物人格中间件 (PetPersonalityMiddleware)                  │
│  思考中→表情 │ 调工具→角色化语言 │ 结果→自然解读            │
│  被拦截→撒娇抱怨 │ 出错→傲娇吐槽 │ 完成→开心表情           │
├──────────────────────────────────────────────────────────┤
│  核心引擎 (CoreEngine)                                     │
│  PreProcessor → SessionEngine → StateLoop → AgentLoop    │
├──────────────────────────────────────────────────────────┤
│  上下文引擎                    │  工具系统                  │
│  SystemPrompt + Memory        │  Skill → MCP/Local/HTTP  │
│  + 附件 + 摘要压缩             │  统一 ToolRegistry        │
│  + 思考强度决策                │  (按模式动态注册)          │
├──────────────────────────────────────────────────────────┤
│  安全控制 (SafetyControl) —— 所有工具调用强制经过           │
│  轻量模式: 简洁拒绝 │ 助手模式: 四级+确认UI+参数级检测       │
├──────────────────────────────────────────────────────────┤
│  平台桥接 (Rust) —— 窗口 / 系统调用 / 后台轮询             │
└──────────────────────────────────────────────────────────┘
''')

add_body('各层职责：')
add_body('• UI 层：Vue 3 组件，负责角色动画渲染、聊天面板展示、设置面板交互')
add_body('• 宠物人格中间件：横切关注点，包裹 Agent 循环的每个阶段，将技术操作翻译为角色化表达')
add_body('• 核心引擎：PreProcessor(消息预处理)、SessionEngine(会话管理)、StateLoop(状态机)、AgentLoop(工具调用循环编排)')
add_body('• 上下文引擎：动态组装 SystemPrompt，按模式/任务类型注入工具声明，管理记忆注入')
add_body('• 工具系统：统一 ToolRegistry，Local/MCP/Skill 三种来源平权注册，ToolRouter 按来源分派执行')
add_body('• 安全控制：所有工具调用强制经过，轻量模式简单拒绝非安全操作，助手模式四级+确认UI')
add_body('• 平台桥接：Rust 提供窗口管理、系统调用(std::process::Command)、MCP stdio 进程桥接')

# ===================
# 第五章
# ===================
doc.add_heading('五、核心 Agent Loop', level=1)

add_diagram_box('Agent Loop 流程图', '''
OriginMessage (用户输入 / 窗口主动)
  │
  ▼
PreProcessor → /slash? 执行返回 │ 空/重复 过滤
  │
  ▼
人格中间件.wrap("thinking") → 表情 + 角色化提示
  │
  ▼
思考强度决策 (thinkingEffort: auto → 分析任务 → low/med/high)
  │
  ▼
ContextEngine.build()
  ├── 1. 人格 Prompt (card + boundary)
  ├── 2. Memory 注入 (相关记忆 topK)
  ├── 3. 工具声明 (按模式 + 任务类型)
  ├── 4. 会话摘要 (如果有)
  └── 5. 输出约束
  │
  ▼
[Plan] 可选 (助手模式 + 复杂任务)
  │
  ▼
Agent Loop 核心循环 ────────────────────────┐
  ├── 1. 请求 AI Model (含思考强度)           │
  ├── 2. 解析输出 (function_call? 纯文本?)     │
  ├── 3. 有工具调用?                           │
  │     → 人格中间件.wrap("executing")         │
  │     → SafetyControl.check()               │
  │     → ToolEngine.execute()                │
  │     → 结果回注 → 回到步骤1                 │
  └── 达上限? → 强制生成回复                   │
  └──────────────────────────────────────────┘
  │
  ▼
人格中间件.wrap("done") → 表情 + 音效 + 边界联动
  │
  ▼
回复生成器 (后处理: kaomoji/截断/HTML转义)
  │
  ▼
UI 展示 (ChatPanel + StreamView + 音效)
''')

doc.add_heading('5.1 Loop 控制参数', level=2)
add_table(
    ['参数', '默认值', '说明'],
    [
        ['maxRetry', '3', '工具调用/验证失败最大重试次数'],
        ['maxToolCallsPerTurn', '5', '单轮最多工具调用链长度'],
        ['toolTimeoutMs', '30000', '单个工具执行超时(ms)'],
        ['turnTimeoutMs', '120000', '整轮总超时(ms)'],
        ['streamEnabled', 'true', '是否流式输出'],
        ['contextCompactAt', '0.95', '上下文利用率达95%触发摘要压缩'],
    ]
)

doc.add_heading('5.2 循环内并行追踪', level=2)
add_body('• usage：token 消耗统计 + 上下文利用率监控')
add_body('• retry count：当前轮次重试计数，达上限强制退出')
add_body('• timeout：整轮超时计时器，超时后停止工具调用')
add_body('• state：Agent 状态机转换追踪 (WAITING → PRE → GENERATING → EXECUTING → ...)')

# ===================
# 第六章
# ===================
doc.add_heading('六、宠物人格中间件', level=1)

add_body('人格不只是 Prompt 生成器，而是包裹整个 Agent 循环的横切中间件。用户永远看到角色化的表达，技术细节（工具调用、安全检查、错误信息）全部在中间件层转换。')

doc.add_heading('6.1 中间件接口', level=2)
add_code('''interface PetPersonalityMiddleware {
  wrap(stage: AgentStage, context: StageContext): PersonalityEffect
  getSystemPrompt(state: PersonalityState): string
  getExpression(stage: AgentStage, state: PersonalityState): string
  getSound(stage: AgentStage, state: PersonalityState): string
}

type AgentStage =
  | "thinking"   // 收到消息，准备思考
  | "planning"   // Plan 步骤进行中
  | "generating" // AI 生成中
  | "executing"  // 执行工具
  | "blocked"    // 安全拦截
  | "error"      // 出错
  | "done"       // 完成回复
  | "idle"       // 空闲

interface PersonalityEffect {
  userMessage?: string     // 展示给用户的角色化提示
  expression: string       // 角色表情 id
  sound?: string           // 音效 id
  boundaryDelta?: number   // 边界值变化
}''')

doc.add_heading('6.2 各阶段行为映射', level=2)
add_table(
    ['阶段', '表情', '用户看到', '音效'],
    [
        ['thinking', 'thinking (思考眨眼)', '无 (除非 > 3秒)', '—'],
        ['planning', 'serious (认真)', '"嗯...让我想想怎么帮你～"', '—'],
        ['generating', 'talk / smile', '流式文本逐字出现', '—'],
        ['executing', 'focus (专注)', '"诶，让我帮你看看..."', '—'],
        ['blocked', 'worried / angry', '"唔...这个我不能做呢～"', '—'],
        ['error', 'cry (哭脸)', '"啊...信号不太好～"', '—'],
        ['done', 'happy / chu', '—', 'reply_ding'],
        ['idle', 'idle (待机)', '主动搭话 (按 boundary 决定)', '—'],
    ]
)

doc.add_heading('6.3 工具 → 角色化文案映射', level=2)
add_table(
    ['工具 ID', '执行中', '完成'],
    [
        ['file.read', '"让我读读这个文件..."', '"读完啦～"'],
        ['file.list', '"让我看看这里面有什么..."', '"看到了～"'],
        ['file.search', '"帮你找找..."', '"找到了！"'],
        ['file.write', '"帮你写下来..."', '"写好啦～"'],
        ['bash.exec', '"让我来处理..."', '"搞定！"'],
        ['system.info', '"检查一下电脑状态..."', '"嗯嗯了解了～"'],
        ['http.get', '"帮你上网查查..."', '"查到了～"'],
        ['app.open', '"帮你打开..."', '"打开了～"'],
        ['MCP 工具 (通用)', '"让我用工具处理一下..."', '"完成啦～"'],
    ]
)

# ===================
# 第七章
# ===================
doc.add_heading('七、工具生态系统', level=1)

doc.add_heading('7.1 统一 ToolRegistry', level=2)
add_body('所有工具（Local / MCP / Skill）统一注册到 ToolRegistry。模型不区分工具来源，只看工具描述来选择。ToolRouter 在背后按 source 字段路由到不同执行器。')

add_diagram_box('工具路由架构', '''
              ToolRegistry（统一注册表）
              ┌───────────────────────────────────────────┐
              │  ToolDef[]  统一接口                        │
              │  { id, name, description, parameters,     │
              │    safetyLevel, handler }                  │
              └───────────────────┬───────────────────────┘
                                  │
              ToolRouter（路由层，按 source 分派）
              ┌───────────────────┼───────────────────────┐
              │                   │                       │
              ▼                   ▼                       ▼
        ┌──────────┐      ┌──────────────┐      ┌──────────────┐
        │  Local   │      │     MCP      │      │    Skill     │
        │  Tools   │      │   Servers    │      │   (编排式)    │
        │          │      │              │      │              │
        │ 始终加载  │      │ 助手模式加载  │      │ 助手模式加载  │
        │          │      │ 工具自动发现  │      │ skills/ 目录  │
        └──────────┘      └──────────────┘      └──────────────┘
''')

doc.add_heading('7.2 ToolDef 标准接口', level=2)
add_code('''interface ToolDef {
  id: string                      // 全局唯一，如 "bash.exec"
  name: string                    // 给模型的函数名
  description: string             // 模型看的描述（决定何时调用）
  parameters: {                   // JSON Schema
    type: "object"
    properties: Record<string, any>
    required: string[]
  }
  safetyLevel: SafetyLevel        // SAFE | NORMAL | DANGER | NOWAY
  source: "local" | "mcp" | "skill"
  sourceId: string                // mcp server id / skill id / "builtin"
  mode: "pet" | "assistant"       // 哪个模式可用
  handler: (params: any, ctx: ToolContext) => Promise<ToolResult>
  timeoutMs?: number
  personalityHint?: {             // 人格化文案模板
    executing?: string
    done?: string
    blocked?: string
  }
}''')

doc.add_heading('7.3 轻量模式工具（始终加载）', level=2)
add_body('轻量模式具备基本的文件、系统、网络能力，能完成日常查询类任务。共 6 个工具，声明简洁（约 150 tokens），始终携带不增加决策开销。')

add_table(
    ['ID', '名称', '安全级别', '说明'],
    [
        ['file.read', '读文件', 'SAFE', '读取文本文件内容'],
        ['file.list', '列目录', 'SAFE', '列出目录内容'],
        ['file.search', '搜索文件', 'SAFE', '按名称/内容搜索'],
        ['system.info', '系统信息', 'SAFE', 'OS/CPU/内存状态'],
        ['bash.exec', '执行命令', 'NORMAL', '仅白名单命令'],
        ['http.get', '网页获取', 'NORMAL', 'GET 请求，返回文本'],
    ]
)

add_body('Bash 白名单（轻量模式下可执行的命令）：ls, cat, head, tail, grep, find, which, echo, pwd, date, whoami, uname, df, du, ps, ifconfig, system_profiler(macOS), tasklist(Win)')

doc.add_heading('7.4 助手模式额外工具', level=2)
add_table(
    ['ID', '名称', '安全级别', '说明'],
    [
        ['file.write', '写文件', 'DANGER', '写入/修改文件'],
        ['file.delete', '删除文件', 'NOWAY', '默认禁止'],
        ['bash.exec_full', '全量命令', 'DANGER', '覆盖白名单限制'],
        ['app.open', '打开应用', 'NORMAL', '启动外部应用'],
        ['clipboard.read', '读剪贴板', 'NORMAL', '读取剪贴板内容'],
        ['clipboard.write', '写剪贴板', 'DANGER', '写入剪贴板'],
        ['agent.spawn', '子代理', 'NORMAL', 'fork/team 模式'],
    ]
)

doc.add_heading('7.5 Skill 系统', level=2)
add_body('Skill 是写在 skills/ 目录下的声明文件（.md），YAML frontmatter + Markdown 执行模板。助手模式启动时自动扫描加载，可热更新(HMR)。')
add_body('Skill 生命周期：Skill 文件 → SkillLoader 解析 → SkillRegistry 注册 → 转为 ToolDef 注入 ToolRegistry → 模型可直接调用 → Skill 执行时内循环调用 Local/MCP 工具 → 拼结果返回')

doc.add_heading('7.6 MCP 集成（助手模式）', level=2)
add_body('MCP (Model Context Protocol) 是工具扩展的标准协议。助手模式下，MCP Manager 读取 CONFIG.yaml 中配置的 MCP Servers，自动连接并 discover 工具，注入 ToolRegistry。')

add_body('支持两种 Transport：')
add_body('• stdio：Rust spawn 子进程，通过 stdin/stdout JSON-RPC 通信。TS 端通过 invoke("mcp_send") 桥接')
add_body('• SSE：TS 端直接 fetch() HTTP SSE 连接，无需 Rust 桥接')

add_body('MCP 生命周期：启动时逐个连接 enabled servers → initialize handshake → tools/list 获取工具列表 → 每个工具转为 ToolDef 注入 Registry。运行时断开自动重连(指数退避，最多3次)。关闭时断开连接并 kill stdio 子进程。')

# ===================
# 第八章
# ===================
doc.add_heading('八、上下文引擎', level=1)

doc.add_heading('8.1 SystemPrompt 动态组装', level=2)
add_body('每次请求前动态组装，不缓存。组装顺序如下：')

add_table(
    ['优先级', '组件', 'Token 估算', '说明'],
    [
        ['1', '人格卡 System Prompt', '~500-2000', '必须，固定'],
        ['2', '边界提示 (boundary)', '~50', 'unansweredCount > 0 时'],
        ['3', 'Memory 注入', '~200-800', 'getRelevant(topK=3)'],
        ['4', '可用工具声明', '~150-1500', '轻量模式 ~150，助手模式按需'],
        ['5', '会话摘要', '~300-500', '上下文压缩后才有'],
        ['6', '输出格式约束', '~50', '固定'],
        ['7', '思考强度提示', '~30', 'low 模式追加"请简要回答"'],
    ]
)

doc.add_heading('8.2 工具声明策略', level=2)
add_body('轻量模式：只有 6 个工具，声明很短（约 150 tokens），始终携带，省去决策逻辑。')
add_body('助手模式：工具可能很多（MCP Server 可能贡献数十个工具），按任务类型分级注入：')
add_body('• L0 (闲聊)：不注入工具声明 — 0 tokens')
add_body('• L1 (可能用工具)：只注入工具名 + 一句话描述 — ~100-200 tokens')
add_body('• L2 (明确工具任务)：全量注入名称 + 参数 schema + 描述 — ~500-1500 tokens')

add_body('检测关键词："帮我" "查看" "打开" "搜索" "找" "整理" "分析" "检查" "文件" "文件夹" "桌面" "下载" "代码" "项目" "系统" "运行" "执行" "命令" "天气" "时间" "日期"')

# ===================
# 第九章
# ===================
doc.add_heading('九、安全控制系统', level=1)

doc.add_heading('9.1 轻量模式安全', level=2)
add_body('简洁策略：SAFE 直接放行，NORMAL / DANGER / NOWAY 直接拒绝（不弹确认 UI）。拒绝时经人格中间件包装为友好提示："这个功能需要在助手模式下使用哦～"')

doc.add_heading('9.2 助手模式安全', level=2)
add_body('完整四级安全检查 + 确认 UI：')

add_table(
    ['级别', '标识', '说明', '示例操作', '行为'],
    [
        ['SAFE', '🟢', '只读/查询', '读文件、查内存、获取系统信息', '自动放行'],
        ['NORMAL', '🟡', '需授权', '打开应用、读剪贴板、发 HTTP GET', '首次确认后可会话内信任'],
        ['DANGER', '🟠', '每次确认', '删改文件、git 操作、执行全量命令', '每次都要用户确认'],
        ['NOWAY', '🔴', '硬禁止', 'sudo、格式化、rm -rf、SSH、密码', '直接拒绝'],
    ]
)

doc.add_heading('9.3 参数级危险检测', level=2)
add_body('在安全检查之前，先对工具调用的参数做模式匹配。匹配到危险模式直接拒绝，不需要走到级别判断。这能防止模型注入恶意参数。')

add_body('Bash 危险模式：rm -rf, sudo, chmod 777, mkfs, dd if=, > /dev/, curl | bash, wget | sh')
add_body('File 危险模式：~/.ssh/, /etc/passwd, /etc/shadow, /System/, C:\\Windows, .pem, .key, .env')

doc.add_heading('9.4 安全确认 UI（人格化）', level=2)
add_body('助手模式下，DANGER 级别的工具调用需要用户确认。确认弹窗使用人格化语言：')
add_code('''┌─────────────────────────────────┐
│  (◕‿◕) 主人～我想执行这个...     │
│                                 │
│  让我看看你的桌面有什么文件～      │
│  可以吗？                       │
│                                 │
│  [ 不行 ]  [ 好的 ]  [ 今天都信你 ] │
└─────────────────────────────────┘''')

# ===================
# 第十章
# ===================
doc.add_heading('十、状态机与记忆系统', level=1)

doc.add_heading('10.1 Agent 状态机', level=2)
add_diagram_box('Agent 状态转换图', '''
WAITING ──(收到消息)──→ PRE ──→ PLANNING(可选)──→ GENERATING
   ▲                       │                          │
   │                       └──────(简单对话)───────────┘
   │                                                    │
   │              ┌─────────────────────────────────────┤
   │              │                                     │
   │              ▼                                     ▼
   └──(完成)── WAITING              ┌──────────── EXECUTING
                                    │                 │
                                    └──(工具完成)──────┘
''')

add_table(
    ['当前状态', '事件', '下一状态'],
    [
        ['WAITING', '收到消息', 'PRE'],
        ['PRE', '预处理完成 / 简单消息', 'GENERATING'],
        ['PRE', '预处理完成 / 复杂任务', 'PLANNING (可选)'],
        ['PLANNING', '规划完成', 'GENERATING'],
        ['GENERATING', '模型返回工具调用', 'EXECUTING'],
        ['GENERATING', '模型返回纯文本', 'WAITING (完成)'],
        ['EXECUTING', '工具执行完成 / 还需工具', 'GENERATING'],
        ['EXECUTING', '达到上限 / 超时', 'WAITING (强制回复)'],
        ['任意', '错误 / 异常', 'WAITING (错误处理)'],
    ]
)

doc.add_heading('10.2 人格状态（三维驱动）', level=2)
add_table(
    ['维度', '0', '1', '2', '≥3'],
    [
        ['unansweredCount', '甜蜜女友', '轻微不满', '占有欲上升', '病娇'],
        ['toolUseCount', '小心翼翼', '逐渐自信', '游刃有余', '—'],
        ['taskComplexity', '轻松(简单)', '专注(中等)', '认真(复杂)', '—'],
    ]
)

doc.add_heading('10.3 记忆系统', level=2)
add_body('长期记忆采用文件注册表模式（与 CLAUDE.md memory 系统一致）：')
add_code('''项目根/memory/
├── MEMORY.md          ← 注册表索引（所有记忆的指针）
├── CANDY.md           ← 用户手写系统指令
├── User.md            ← 用户画像、偏好、习惯
├── Outside.md         ← 外部知识指针（在哪找什么）
└── Project/           ← 按 SessionID 归档的会话上下文''')

add_body('短期记忆：会话级，包含原始对话（最近 maxContextMessages 条）、压缩摘要（95% 上下文时触发全量摘要）、工作记忆指针。')

add_body('记忆整理（两个模式共用，实现方式不同）：触发条件为 24h 定时器 / 2 个会话结束 / 手动 /memory clean。轻量模式使用 JS 定时器触发，主线程同步执行简单扫描去重，不阻塞 UI。助手模式使用后台 fork agent 执行，含 AI 辅助判断矛盾、冲突标记，更深度。核心流程统一：扫描记忆文件 → 去重(相似度>0.8) → 过期移除(30天+低重要性) → 矛盾标记 → 重写 MEMORY.md 注册表 → 释放 .lock 文件锁。')

# ===================
# 第十一章
# ===================
doc.add_heading('十一、启动装配流程', level=1)

add_diagram_box('应用启动流程', '''
main.ts
  │
  ├── 1. Rust 后端初始化
  │   ├── ActivationPolicy::Accessory (macOS 隐藏 Dock)
  │   ├── create_main_window (透明置顶悬浮)
  │   ├── 系统托盘创建
  │   └── 后台监控线程启动 (窗口标题捕获)
  │
  ├── 2. 共用模块初始化 (始终加载)
  │   ├── Logger (读 loggingConfig.level)
  │   ├── Config 加载 (CONFIG.yaml → 类型化 getter)
  │   ├── Personality 注册表 + 人格中间件
  │   ├── MemorySystem (读 MEMORY.md 注册表)
  │   ├── ToolRegistry (轻量模式 6 个工具)
  │   ├── SafetyControl (简洁模式)
  │   ├── Provider (AI HTTP)
  │   └── ContextEngine / ReplyGenerator
  │
  ├── 3. if (assistantMode):
  │   ├── dynamic import → local-extra/ mcp/ skill/ safety-full
  │   ├── 追加注册助手专属工具
  │   ├── 初始化 MCP Manager (异步连接，不阻塞 UI)
  │   ├── 加载 Skill 目录
  │   └── Safety 升级为完整模式
  │
  ├── 4. 恢复上次会话
  │   ├── 读 localStorage → deskpet_session_current
  │   ├── 有上次会话 → resumeSession()
  │   └── 无 → createSession() → initChat(欢迎消息)
  │
  ├── 5. 加载 CANDY.md + 长期记忆 → 注入当前会话
  │
  └── 6. UI 就绪 → WAITING 状态
''')

# ===================
# 第十二章
# ===================
doc.add_heading('十二、关键调用链', level=1)

doc.add_heading('12.1 轻量模式：日常文件查询', level=2)
add_diagram_box('', '''
用户: "桌面有什么文件"
  → PreProcessor → 非命令
  → 人格中间件.wrap("thinking") → expression: "thinking"
  → 思考强度: auto → 检测"什么/文件" → medium
  → ContextEngine.build()
     ├── 人格 Prompt (KAngel 甜蜜女友 + boundary=2)
     ├── 轻量工具声明 (6个, ~150 tokens)
     └── Memory 注入 (相关记忆 topK=3)
  → AI → tool_call: file.list({ path: "~/Desktop" })
  → 人格中间件.wrap("executing") → "让我看看桌面有什么..."
  → SafetyControl.check("file.list", SAFE) → 放行
  → Rust read_dir → { entries: [...] }
  → 人格中间件.wrap("result") → "嗯看到了"
  → 结果回注 → AI 生成自然语言
  → "桌面上有3个文件夹和2个文件呢～Documents、Pictures..."
  → 人格中间件.wrap("done") → expression: "happy", sound: "reply"
  → ChatPanel 流式显示 + StreamView 表情切换 + 音效
  → StateLoop: WAITING (unansweredCount reset)
''')

doc.add_heading('12.2 助手模式：复杂工具链', level=2)
add_diagram_box('', '''
用户: "帮我整理桌面，把截图放到一个文件夹"
  → PreProcessor → 检测 "帮我/整理/截图/文件夹"
  → 人格中间件.wrap("thinking")
  → 思考强度: auto → 检测多关键词 → high
  → ContextEngine.build() → 助手模式全量工具声明
  → Plan: 1.列文件 2.找截图 3.创建文件夹 4.移动文件

  → [Round 1] AI → file.list("~/Desktop")
     → SAFE → 放行 → 返回文件列表

  → [Round 2] AI → bash.exec("find ~/Desktop -name '*.png' -o -name '*.jpg'")
     → DANGER → 确认UI(人格化) → 用户点击"好的" → 执行
     → 返回 a.png, b.jpg

  → [Round 3] AI → bash.exec("mkdir -p ~/Desktop/截图整理")
     → DANGER → 确认 → 执行 → 成功

  → [Round 4] AI → bash.exec("mv ~/Desktop/*.png ~/Desktop/*.jpg ~/Desktop/截图整理/")
     → DANGER → 确认 → 执行 → 成功

  → 最终回复: "搞定啦～已经把截图整理到'截图整理'文件夹了♡"
  → 人格中间件.wrap("done") → expression: "chu"
''')

doc.add_heading('12.3 MCP 工具调用（助手模式）', level=2)
add_diagram_box('', '''
用户: "数据库里有什么表"
  → MCP Server "sqlite" 已连接，tools 已注入 Registry
  → ContextEngine → 包含 MCP 工具声明
     "mcp.sqlite.list_tables: 列出数据库所有表"
     "mcp.sqlite.describe_table: 查看表结构"

  → AI → tool_call: mcp.sqlite.list_tables({})
     → 人格中间件.wrap("executing") → "让我查查数据库..."
     → SafetyControl: NORMAL → 首次确认 → 用户"好的"
     → MCP Client → tools/call JSON-RPC → sqlite MCP server
     → 返回: ["users", "orders", "products"]

  → AI → tool_call: mcp.sqlite.describe_table({ table: "users" })
     → NORMAL → 会话已信任 → 直接放行
     → 返回: [{ name: "id", type: "INTEGER" }, ...]

  → AI 最终回复: "这个数据库有3个表：users(用户表)、orders(订单表)、products(商品表)。users表有id、name、email字段..."
''')

# ===================
# 第十三章
# ===================
doc.add_heading('十三、目录结构', level=1)

add_code('''Desk-Pet/
├── CONFIG.yaml / CONFIG-DEV.yaml    # 全局配置
├── DES.md / CLAUDE.md               # 设计/开发文档
├── DESIGN_ORIGIN.md                 # 设计草案
├── 架构方案.md                       # 架构设计文档
│
├── memory/                          # 长期记忆（文件注册表）
│   ├── MEMORY.md / CANDY.md
│   ├── User.md / Outside.md
│   └── Project/
│
├── sessions/                        # 历史会话持久化
├── skills/                          # Skill 文件目录 (助手模式)
│
├── src/                             # Vue 3 + TypeScript
│   ├── App.vue / main.ts
│   └── services/
│       ├── engine/                  # 核心引擎
│       │   ├── preprocessor.ts      # Slash命令 + 路由
│       │   ├── session.ts           # 会话引擎
│       │   ├── state-loop.ts        # 状态机
│       │   ├── agent-loop.ts        # ★ 唯一 Agent Loop
│       │   ├── thinking.ts          # ★ 思考强度决策
│       │   └── parser.ts            # AI输出解析
│       │
│       ├── personality/             # 人格模块
│       │   ├── middleware.ts        # ★ 人格中间件
│       │   ├── registry.ts / loader.ts
│       │   ├── boundary.ts / types.ts
│       │   └── cards/
│       │
│       ├── tool/                    # 工具系统
│       │   ├── registry.ts          # ★ 统一注册表 (按模式)
│       │   ├── router.ts            # 工具路由
│       │   ├── local/               # 始终加载
│       │   │   ├── file.ts / bash.ts
│       │   │   ├── system.ts / http.ts
│       │   ├── local-extra/         # ★ 助手模式动态加载
│       │   │   ├── file-write.ts / bash-full.ts
│       │   │   ├── app.ts / clipboard.ts
│       │   │   └── agent-tool.ts
│       │   ├── skill/               # ★ 助手模式
│       │   │   ├── loader.ts / registry.ts / runner.ts
│       │   ├── mcp/                 # ★ 助手模式
│       │   │   ├── manager.ts / client.ts
│       │   │   ├── stdio.ts / sse.ts
│       │   └── http/
│       │
│       ├── safety/                  # 安全控制
│       │   ├── checker.ts           # 基础检查 (始终)
│       │   └── full-checker.ts      # ★ 完整检查 (助手模式)
│       │
│       ├── context/                 # 上下文引擎
│       │   ├── builder.ts           # SystemPrompt 组装
│       │   └── tool-selector.ts     # 工具声明决策
│       │
│       ├── memory/                  # 记忆系统
│       ├── provider/                # AI Provider
│       ├── reply/                   # 回复生成器
│       ├── window/ / audio/
│       ├── config.ts / logger.ts / cooldown.ts / env.ts
│       └── ...
│
├── src-tauri/                       # Rust 后端
│   └── src/
│       ├── lib.rs / monitor/ / window/
│       └── commands/
│           ├── cursor.rs / monitor_ctl.rs / sim.rs / logging.rs
│           ├── tool_exec.rs         # ★ Bash/文件执行
│           └── mcp_bridge.rs        # ★ MCP stdio 桥接
│
└── public/assets/''')

# ===================
# 第十四章
# ===================
doc.add_heading('十四、配置系统', level=1)

add_body('以下为在现有 CONFIG.yaml 基础上新增的配置段。所有配置项可在设置面板编辑，部分即时生效，部分重启后生效。')

add_code('''# ── 模式 ──
mode:
  assistant: false                # 助手模式开关 (默认轻量)

# ── AI 思考强度 ──
ai:
  thinkingEffort: "auto"          # auto | low | medium | high
  thinkingBudget:
    low: 1000                     # low 模式最大 thinking tokens
    medium: 4000
    high: 16000

# ── 核心 Loop ──
loop:
  maxRetry: 3                     # 最大重试次数
  maxToolCallsPerTurn: 5          # 单轮最多工具调用链
  toolTimeoutMs: 30000            # 单个工具超时(ms)
  turnTimeoutMs: 120000           # 整轮超时(ms)
  streamEnabled: true             # 是否流式输出
  contextCompactAt: 0.95          # 上下文压缩阈值

# ── 工具系统 ──
tools:
  bash:
    enabled: true
    whitelist: ["ls", "cat", "head", "tail", "grep", "find",
                "which", "echo", "pwd", "date", "whoami",
                "uname", "df", "du", "ps", "tasklist",
                "system_profiler", "ifconfig"]
  file:
    enabled: true
    writeEnabled: false           # 轻量模式强制 false
  mcp:
    enabled: false                # 助手模式才有效
    defaultSafetyLevel: "normal"
    connectionTimeoutMs: 10000
    reconnectMaxRetries: 3
    servers: []                   # MCP Server 配置列表
  skill:
    enabled: false                # 助手模式才有效
    directory: "skills/"

# ── 安全 ──
safety:
  mode: "tell_me"                 # just_do_it | tell_me | let_me_tk
  sessionTrustEnabled: true       # 会话内信任 NORMAL 操作
  dangerousParamCheck: true       # 参数级危险检测''')

# ===================
# 第十五章
# ===================
doc.add_heading('十五、性能设计', level=1)

doc.add_heading('15.1 内存预算', level=2)
add_table(
    ['组件', '轻量模式', '助手模式增量'],
    [
        ['WebView + Vue', '~40-60 MB (WebKit 基础开销)', '—'],
        ['JS Heap', '~5-10 MB', '+2-4 MB'],
        ['Rust 进程', '~5-10 MB', '—'],
        ['MCP 子进程', '0', '每进程 +20-50 MB'],
        ['会话数据', '< 2 MB', '—'],
        ['ToolRegistry', '~0.2 MB', '+0.5 MB'],
        ['总计', '~50-80 MB', '~55-150 MB (含MCP)'],
    ]
)

doc.add_heading('15.2 Token 优化', level=2)
add_table(
    ['场景', '轻量模式', '助手模式 L1', '助手模式 L2', '节省'],
    [
        ['简单聊天', '~800 tokens', '~800 tokens', '—', '60% vs 全量'],
        ['日常查询', '~950 tokens', '~1200 tokens', '—', '—'],
        ['复杂任务', '—', '—', '~2000-3000 tokens', '33% vs 全量'],
    ]
)

doc.add_heading('15.3 优化策略', level=2)
add_body('1. 工具声明按需注入 — 轻量模式 6 个工具声明短(~150 tokens)，始终携带；助手模式按任务分级(L0/L1/L2)')
add_body('2. 流式输出 — 边生成边展示，降低感知延迟和 JS Heap 堆积')
add_body('3. 摘要压缩 — 上下文使用率达 95% 时触发全量摘要，将长对话压缩为结构化摘要，节省约 60% token')
add_body('4. 惰性加载 — 助手模式专属模块（MCP/Skill/local-extra/safety-full）通过 dynamic import 按需加载')
add_body('5. 结果截断 — 工具返回值硬限制 50KB，超出截断并标记 truncated: true')
add_body('6. 子代理复用 — fork agent 执行完毕立即销毁上下文，不保留')
add_body('7. 连接池复用 — HTTP fetch 复用 keep-alive 连接')
add_body('8. 模式降级 — 从助手切换回轻量时，释放所有助手模块引用，让 GC 回收')

# ===================
# 第十六章
# ===================
doc.add_heading('十六、接口契约', level=1)

doc.add_heading('16.1 Rust ↔ TS IPC（新增命令）', level=2)
add_code('''// 工具执行
invoke("bash_exec", { command, cwd?, timeout? })
  → { stdout: string, stderr: string, exitCode: number }

invoke("file_read", { path })
  → { content: string, size: number }

invoke("file_write", { path, content })
  → { success: boolean, error?: string }

invoke("file_list", { path })
  → { entries: { name: string, kind: "file"|"dir", size: number }[] }

invoke("system_info")
  → { os: string, arch: string, cpuCount: number, memTotal: number, memUsed: number }

invoke("app_open", { path }) → { success: boolean }
invoke("clipboard_read") → { text: string }
invoke("clipboard_write", { text }) → { success: boolean }

// MCP 桥接 (助手模式)
invoke("mcp_spawn", { command, args, env? }) → { serverId: string, pid: number }
invoke("mcp_send", { serverId, message: string }) → { response: string }
invoke("mcp_kill", { serverId }) → { success: boolean }''')

doc.add_heading('16.2 Events (Rust → TS)', level=2)
add_code('''window-changed   → { title, exe, timestamp }
mcp-status       → { serverId, status: "connected"|"disconnected"|"error", error? }
tool-executing   → { toolId, toolName, personalityHint? }   // UI 状态显示
tool-completed   → { toolId, success, personalityHint? }''')

doc.add_heading('16.3 localStorage Keys', level=2)
add_code('''deskpet_session_current      → 当前会话状态 JSON
deskpet_session_index        → 会话索引 (元数据列表)
deskpet_user_settings        → 用户 UI 设置
deskpet_config_overrides     → 配置覆盖值
deskpet_sound_assignments    → 音效分配
deskpet_unanswered           → 未回复计数
deskpet_safety_trusts        → 会话内信任的工具列表
deskpet_mcp_configs          → MCP Server 配置覆盖 (本地)''')

# ===================
# 第十七章
# ===================
doc.add_heading('十七、迁移路径', level=1)

add_body('从当前架构到目标架构的逐步迁移，每一步不破坏现有功能。')

add_table(
    ['阶段', '内容', '影响范围', '预计周期'],
    [
        ['Phase 1', '人格中间件 + 思考强度决策\n• personality/middleware.ts\n• engine/thinking.ts\n• sendMessage() 插入钩子', 'personality/ engine/ agent/', '1-2 周'],
        ['Phase 2', '工具系统 + 轻量模式\n• tool/registry.ts + router.ts\n• 轻量6工具实现\n• SafetyControl(简洁)\n• 接入AgentLoop\n• Rust tool_exec.rs', 'tool/ safety/ engine/ rust/', '2-3 周'],
        ['Phase 3', '助手模式\n• 模式开关 + 设置UI\n• 助手专属工具\n• 完整Safety + 确认UI\n• agent.spawn 子代理', 'tool/local-extra/ safety/ UI/', '2-3 周'],
        ['Phase 4', 'MCP + Skill\n• MCP Manager/Client (stdio/SSE)\n• Skill Loader/Runner\n• Rust mcp_bridge.rs\n• Plan 步骤', 'tool/mcp/ tool/skill/ rust/', '3-4 周'],
    ]
)

# ===================
# 第十八章
# ===================
doc.add_heading('十八、关键设计决策', level=1)

add_table(
    ['决策', '选择', '理由'],
    [
        ['双模式', '同一 Loop + 不同工具集', '不搞两套代码，共用引擎，零冗余'],
        ['轻量能力', '6 个基础工具', '覆盖日常查询（读文件/列目录/系统信息/安全命令），不残废'],
        ['思考强度', 'auto 自动调节', '闲聊快响应(节省token)，复杂任务深推理(提高准确率)'],
        ['人格定位', '横切中间件', 'executing/blocked/error 全部阶段都要角色化，不只是 Prompt'],
        ['MCP', '助手模式专属', '普通用户不需要外部工具生态'],
        ['Skill', '助手模式专属', '编排式工作流属高级功能，轻量用户不需要'],
        ['工具声明(轻量)', '始终携带', '仅6个工具~150 tokens，不值得做决策逻辑'],
        ['工具声明(助手)', '按需三级注入', '工具可能很多(MCP贡献数十个)，必须分级省 token'],
        ['Rust 职责', '保持轻薄 OS 层', '性能瓶颈在网络不在本地，TS 灵活做 Logic'],
        ['状态管理', 'Vue reactive', '保持 UI 响应式，状态可序列化持久化'],
        ['记忆存储', '文件系统 MD + localStorage', '可读、可迁移、低内存'],
        ['SubAgent', 'TS 内实现', '不需要系统级进程隔离，fork 模式用独立上下文即可'],
    ]
)

# ── Save ──
output_path = '/Users/v1rtual/WEB/someworks/Desk-Pet/架构设计文档.docx'
doc.save(output_path)
print(f'✅ 文档已生成: {output_path}')
